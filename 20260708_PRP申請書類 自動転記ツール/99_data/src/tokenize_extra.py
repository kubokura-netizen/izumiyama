# -*- coding: utf-8 -*-
"""
find一致では拾えない「複数段落ブロック」をトークン化する追加処理。
  ・治療価格ブロック（4.5説明書）… 「①〇〇キットを用いた治療：〇円」の連続段落 → {{治療価格}}
  ・お問合せ先ブロック（4.5説明書）… 「■お問合せ先…」〜「※万一…」 → {{患者からの問い合わせ先}}
テンプレ(マスター)を直接書き換える。エンジンは docx_kit の _fill_token_multiline で穴埋めする。
冪等：既に {{…}} 済みなら対象段落が見つからずスキップ。
"""
import glob, os, sys, re

SRC_DIR = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.dirname(SRC_DIR)
TPL = os.path.join(DATA, "テンプレート")
sys.path.insert(0, SRC_DIR)
from docx import Document
from transcribe_sop import _run_green

PRICE_RE = re.compile(r"を用いた治療[：:]")


def _set_para_token(p, token):
    """段落を単一の緑ラン {token} に置き換える。"""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(token)
    _run_green(run)


def tokenize_price(d):
    paras = d.paragraphs
    idxs = [i for i, p in enumerate(paras) if PRICE_RE.search(p.text) and "{{" not in p.text]
    if not idxs:
        return 0
    first = idxs[0]
    _set_para_token(paras[first], "{{治療価格}}")
    for i in idxs[1:]:
        for r in list(paras[i].runs):
            r._element.getparent().remove(r._element)
    return 1


def tokenize_inquiry(d):
    paras = d.paragraphs
    start = next((i for i, p in enumerate(paras) if "■お問合せ先" in p.text and "{{患者からの問い合わせ先}}" not in p.text), None)
    if start is None:
        return 0
    # ブロック終端：start以降で最初に「※万一」を含む段落
    end = next((i for i in range(start, min(len(paras), start + 15)) if "※万一" in paras[i].text), None)
    if end is None:
        end = start   # 見つからなければ先頭段落のみ
    # 終端段落の「※万一…（ください/下さい）。」より後ろ（例：患者様記入欄）は残す
    tail = ""
    if end != start:
        m = re.search(r"※万一.*?(?:ください|下さい)。", paras[end].text)
        if m:
            tail = paras[end].text[m.end():]
    else:
        # 先頭段落に※万一が同居する場合も末尾を保持
        m = re.search(r"※万一.*?(?:ください|下さい)。", paras[start].text)
        if m:
            tail = paras[start].text[m.end():]
    # start段落 → トークン、中間段落クリア、end段落 → tailのみ
    _set_para_token(paras[start], "{{患者からの問い合わせ先}}")
    for i in range(start + 1, end + 1):
        for r in list(paras[i].runs):
            r._element.getparent().remove(r._element)
    if tail.strip():
        paras[end].add_run(tail)
    return 1


def remove_redundant_suffix(d):
    """採血量の注記がテンプレ固定文（半角括弧）で重複する分を除去（値側が全角で付与するため）。"""
    from transcribe_sop import _iter_all_paragraphs
    DUP = "（使用キットにより異なる)"     # 半角閉じ括弧＝テンプレ固定側
    n = 0
    for p in _iter_all_paragraphs(d):
        for r in p.runs:
            if DUP in r.text:
                r.text = r.text.replace(DUP, "")
                n += 1
    return n


def main():
    total = 0
    for key in ["2種関節系PRP", "3種筋腱靭帯系PRP"]:
        for f in glob.glob(os.path.join(TPL, key, "4.5*.docx")):
            if os.path.basename(f).startswith("~"):
                continue
            d = Document(f)
            a = tokenize_price(d)
            b = tokenize_inquiry(d)
            s = remove_redundant_suffix(d)
            if a or b or s:
                d.save(f)
                total += a + b
                print("[%s] %s: price=%d inquiry=%d suffix_fix=%d" % (key, os.path.basename(f)[:20], a, b, s))
    print("total blocks tokenized:", total)


if __name__ == "__main__":
    main()
