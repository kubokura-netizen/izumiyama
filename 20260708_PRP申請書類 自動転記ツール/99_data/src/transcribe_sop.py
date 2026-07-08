# -*- coding: utf-8 -*-
"""
SOP（Word/.docx）へヒアリングシートの内容を差し込むモジュール。

処理:
  1) 99_data/マッピング/sop_replacements.json を読み込む（定型サンプル値→転記元）。
  2) テンプレフォルダ 99_data/テンプレート/SOP/ を 02_output/SOP_(ヒアリング名)/ へ
     フォルダごとコピー（中のファイル名は変更しない）。
  3) コピー先の各 .docx について、本文・表・ヘッダ・フッタを走査し、定型値を
     ヒアリングの値へ置換する（図・表・書式はそのまま保持）。

※ python-docx を使用（Word不要・.docx対応）。.dotx 等はコピーのみで置換対象外。
※ クリニック名・氏名・制定日が主な対象。設定は sop_replacements.json で変更可能。
"""
import os, io, json, glob, shutil, re


def _clean(s):
    return "" if s is None else str(s).strip()


# ---- パス ----
SRC_DIR = os.path.dirname(os.path.abspath(__file__))   # 99_data/src
DATA = os.path.dirname(SRC_DIR)                         # 99_data
SOP_CONFIG = os.path.join(DATA, "マッピング", "sop_replacements.json")


def _resolve_value(src, hearing, run_dt, find):
    """置換後の値を返す。find文字列に応じて日付の粒度（年月日/月日）を切替える。"""
    t = src.get("t")
    if t == "hearing":
        v = hearing.lookup(src.get("label"), src.get("section", ""), int(src.get("occ", 1)))
        return _clean(v)
    if t == "today":
        y, m, d = run_dt.year, run_dt.month, run_dt.day
        # find に「年」が含まれれば年月日、無ければ月日のみ
        if "年" in find:
            return "%d年%d月%d日" % (y, m, d)
        return "%d月%d日" % (m, d)
    if t == "fixed":
        return _clean(src.get("value"))
    return ""


def _build_pairs(config, hearing, run_dt):
    """[({{トークン}}, repl), …] を生成。トークン方式：テンプレの {{項目名}} を穴埋め。
       値は core の resolve を使い、日付は source の fmt（ymd/slash）どおりに整形する。"""
    from transcribe import token_for, resolve as tx_resolve
    try:
        kits = hearing.prp_kits()
    except Exception:
        kits = []
    pairs = []
    for rep in config.get("replacements", []):
        src = rep.get("source", {})
        val, _ = tx_resolve(src, hearing, kits)
        repl = _clean(val)
        if not repl:
            continue
        tok = token_for(src)
        if tok:
            pairs.append(("{{%s}}" % tok, repl))          # 主：トークン穴埋め
        finds = rep.get("find", [])
        if isinstance(finds, str):
            finds = [finds]
        for find in finds:                                 # 副：サンプル値フォールバック
            if find and repl != find:
                pairs.append((find, repl))
    pairs.sort(key=lambda fr: len(fr[0]), reverse=True)   # 長い順（部分一致の取りこぼし防止）
    return pairs


def _replace_in_paragraph(p, pairs):
    cnt = 0
    # 1) ラン単位（書式を保持したまま置換）
    for run in p.runs:
        t = run.text
        if not t:
            continue
        nt = t
        for find, repl in pairs:
            if find in nt:
                cnt += nt.count(find)
                nt = nt.replace(find, repl)
        if nt != t:
            run.text = nt
            _run_green(run)                 # 転記箇所を緑にマーク
    # 2) ラン跨ぎ（1つのランに分断された場合は先頭ランへ集約）
    full = "".join(r.text for r in p.runs)
    if full and any(find in full for find, _ in pairs):
        new = full
        for find, repl in pairs:
            if find in new:
                cnt += new.count(find)
                new = new.replace(find, repl)
        if p.runs:
            p.runs[0].text = new
            _run_green(p.runs[0])
            for r in p.runs[1:]:
                r.text = ""
    return cnt


# ---- 転記箇所の色マーク（緑=転記済み） ----
def _run_green(run):
    try:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
    except Exception:
        pass


def _is_green_rgb(col):
    try:
        if col is None or getattr(col, "rgb", None) is None:
            return False
        s = str(col.rgb).upper()
        return s.endswith("00B050") or s.endswith("008000")
    except Exception:
        return False


def reset_doc_green_to_black(doc):
    """テンプレに残る緑文字を一旦すべて黒に戻す（後で転記箇所だけ緑に付け直す）。"""
    try:
        from docx.shared import RGBColor
    except Exception:
        return 0
    n = 0
    for p in _iter_all_paragraphs(doc):
        for r in p.runs:
            if _is_green_rgb(r.font.color):
                try:
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    n += 1
                except Exception:
                    pass
    return n


def _replace_lowlevel(doc, pairs):
    """全 <w:t> 要素を直接走査して置換（段落/表APIが届かないテキストボックス等も対象）。
       ラン跨ぎは拾えないが、1要素内で連続する定型値（クリニック名等）を確実に置換する。"""
    try:
        from docx.oxml.ns import qn
    except Exception:
        return 0
    roots = [doc.element]
    for sec in doc.sections:
        for hf in (sec.header, sec.footer,
                   getattr(sec, "first_page_header", None), getattr(sec, "first_page_footer", None)):
            el = getattr(hf, "_element", None)
            if el is not None:
                roots.append(el)
    cnt = 0
    seen = set()
    for root in roots:
        if id(root) in seen:
            continue
        seen.add(id(root))
        for t in root.iter(qn("w:t")):
            if not t.text:
                continue
            nt = t.text
            for find, repl in pairs:
                if find in nt:
                    cnt += nt.count(find)
                    nt = nt.replace(find, repl)
            if nt != t.text:
                t.text = nt
    return cnt


def _fill_token_multiline(doc, token, text):
    """{{トークン}}を含む段落を、複数行テキスト（\\n区切り）で置き換える。
       段落内改行は <w:br/> で表現し、1段落に収める（書式は段落既定を継承）。"""
    n = 0
    lines = (text or "").split("\n")
    for p in _iter_all_paragraphs(doc):
        if token not in p.text:
            continue
        for r in list(p.runs):          # 既存ランを除去
            r._element.getparent().remove(r._element)
        run = p.add_run(lines[0] if lines else "")
        _run_green(run)
        for ln in lines[1:]:
            run.add_break()
            run = p.add_run(ln)
            _run_green(run)
        n += 1
    return n


def _set_cell_multiline(cell, text):
    """表セルの内容を text（\\n区切り）で置き換える。1段落目に集約し改行は <w:br/>。"""
    text = "" if text is None else str(text)
    lines = text.split("\n")
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(lines[0] if lines else "")
    _run_green(run)
    for ln in lines[1:]:
        run.add_break()
        run = p.add_run(ln)
        _run_green(run)
    # 余分な段落は空に
    for extra in cell.paragraphs[1:]:
        for r in list(extra.runs):
            r._element.getparent().remove(r._element)
    return True


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nt in cell.tables:
                for q in _iter_table_paragraphs(nt):
                    yield q


def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for p in _iter_table_paragraphs(t):
            yield p
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
            try:
                for p in part.paragraphs:
                    yield p
                for t in part.tables:
                    for q in _iter_table_paragraphs(t):
                        yield q
            except Exception:
                pass


_SOP_TOKEN_RE = re.compile(r"\{\{[^{}]*\}\}")


def _clear_sop_tokens(doc):
    """穴埋めされずに残った {{...}} を空欄化（未充足トークンを見せない）。"""
    n = 0
    for p in _iter_all_paragraphs(doc):
        for r in p.runs:
            if r.text and "{{" in r.text and "}}" in r.text:
                nt = _SOP_TOKEN_RE.sub("", r.text)
                if nt != r.text:
                    r.text = nt
                    n += 1
        full = "".join(r.text for r in p.runs)
        if "{{" in full and "}}" in full and p.runs:
            new = _SOP_TOKEN_RE.sub("", full)
            if new != full:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
                n += 1
    return n


def run_sop(hearing, hearing_path, dir_tpl, dir_output, run_log, run_dt):
    """SOPフォルダを差し込み出力する。戻り値: 出力フォルダパス or None。"""
    try:
        from docx import Document
    except ImportError:
        run_log.append("[SOP] python-docx 未導入のためSOPはスキップ（pip install python-docx）")
        return None
    if not os.path.exists(SOP_CONFIG):
        run_log.append("[SOP] sop_replacements.json が無いためスキップ")
        return None
    with io.open(SOP_CONFIG, encoding="utf-8") as f:
        config = json.load(f)

    tpl_folder = os.path.join(dir_tpl, config.get("template_folder", "SOP"))
    if not os.path.isdir(tpl_folder):
        run_log.append("[SOP] テンプレフォルダ無し: %s" % tpl_folder)
        return None

    base = os.path.splitext(os.path.basename(hearing_path))[0]
    out_folder = os.path.join(dir_output, config.get("output_prefix", "SOP_") + base)

    # フォルダごとコピー（中のファイル名は変更しない）
    if os.path.exists(out_folder):
        shutil.rmtree(out_folder)
    shutil.copytree(tpl_folder, out_folder)

    pairs = _build_pairs(config, hearing, run_dt)
    if not pairs:
        run_log.append("[SOP] 置換対象の値がヒアリングから取得できずスキップ（フォルダはコピー済）")
        return out_folder

    total, n_files = 0, 0
    for f in sorted(glob.glob(os.path.join(out_folder, "*.docx"))):
        if os.path.basename(f).startswith("~$"):
            continue
        try:
            doc = Document(f)
        except Exception as e:
            run_log.append("[SOP] 読込失敗 %s: %r" % (os.path.basename(f), e))
            continue
        cnt = 0
        for p in _iter_all_paragraphs(doc):
            cnt += _replace_in_paragraph(p, pairs)
        leftover = _clear_sop_tokens(doc)
        if cnt or leftover:
            doc.save(f)
            n_files += 1
            total += cnt
    run_log.append("[SOP] 出力フォルダ: %s（%d/%d ファイルに %d件差し込み）"
                   % (os.path.basename(out_folder),
                      n_files, len(glob.glob(os.path.join(out_folder, "*.docx"))), total))
    run_log.append("[SOP] 差し込み内容: " +
                   " / ".join("%s→%s" % (f, r) for f, r in pairs))
    return out_folder


if __name__ == "__main__":
    # 単体実行用（transcribe.py の Hearing を流用）
    import sys, datetime
    from transcribe import Hearing, find_hearing, resolve_dir, DIR_TPL, DIR_OUTPUT
    hp = find_hearing()
    hearing = Hearing(hp, "ヒアリングシート（PRP）")
    log = []
    out = run_sop(hearing, hp, DIR_TPL, DIR_OUTPUT, log, datetime.datetime.now())
    print("=== SOP 差し込み完了 ===")
    print("入力 :", hp)
    print("出力 :", out)
    for line in log:
        print("  -", line)
