# -*- coding: utf-8 -*-
"""
フォルダ型テンプレ（2種関節系PRP / 3種筋腱靭帯系PRP / SOP）への転記エンジン。

各テンプレは「フォルダ」で、中に
  ・様式第一の二（.xlsx）… セル転記
  ・その他の書類（.docx）… サンプル値の一括置換（find→ヒアリング値）
  ・PDF・論文・追加xlsx 等 … 無変更で同梱（参照用）
を含む。

処理:
  1) 99_data/マッピング/docs_config.json を読み込む。
  2) テンプレフォルダを 02_output/(フォルダ名)_(ヒアリング名)/ へフォルダごとコピー。
  3) 様式xlsx はセル転記（resolve→安全書込）。
  4) 直下の各 .docx はサンプル値を一括置換（本文・表・ヘッダ/フッタ、図・書式は保持）。
  5) サブフォルダ内（論文PDF等）は対象外＝無変更で同梱。

※ Excel処理は transcribe.py（openpyxl）、Word処理は transcribe_sop.py の置換関数を再利用。
"""
import os, io, json, glob, shutil

SRC_DIR = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.dirname(SRC_DIR)
DOCS_CONFIG = os.path.join(DATA, "マッピング", "docs_config.json")
FURIGANA_MAP = os.path.join(DATA, "マッピング", "ふりがな.json")


def _load_furigana():
    """医療機関名→ルビ（ふりがな）の手入力マップを読む。無ければ空。"""
    try:
        d = json.load(io.open(FURIGANA_MAP, encoding="utf-8"))
        return d.get("読み", {}) if isinstance(d, dict) else {}
    except Exception:
        return {}


def _clean(s):
    return "" if s is None else str(s).strip()


def _long(path):
    """Windowsの260文字パス上限を回避する拡張パス（\\\\?\\）へ変換。
       深いサブフォルダ内の長いPDF名で copytree が WinError 3 になるのを防ぐ。
       他OSではそのまま返す。"""
    if os.name != "nt":
        return path
    ap = os.path.abspath(path)
    if ap.startswith("\\\\?\\"):
        return ap
    if ap.startswith("\\\\"):                     # UNCパス
        return "\\\\?\\UNC\\" + ap[2:]
    return "\\\\?\\" + ap


# ---- Excelセルの色マーク（緑=転記済み） ----
def _is_green_cell(cell):
    try:
        col = cell.font.color
        if col is not None and getattr(col, "rgb", None):
            s = str(col.rgb).upper()
            return s.endswith("00B050") or s.endswith("008000")
    except Exception:
        pass
    return False


def _set_cell_color(cell, hexcolor):
    from openpyxl.styles import Font
    f = cell.font
    cell.font = Font(name=f.name, size=f.size, bold=f.bold, italic=f.italic,
                     underline=f.underline, strike=f.strike, vertAlign=f.vertAlign,
                     color=hexcolor)


GREEN = "FF00B050"
BLACK = "FF000000"


def _build_docx_pairs(entries, hearing, kits, TX):
    """docx置換ペアを生成（空値は除外、長い順）。ハイブリッド方式：
       ・主：テンプレの {{ヒアリング項目名}} を実値へ穴埋め（トークン方式）。
       ・副：旧サンプル値(find)→実値の置換も併載（トークン化漏れの緑を取りこぼさない）。
       トークン化済みの緑は {{...}} なのでサンプルfindは発火せず二重置換にならない。"""
    pairs = []
    for rep in entries:
        src = rep.get("source", {})
        val, _ = TX.resolve(src, hearing, kits)
        val = _clean(val)
        if not val:
            continue
        fc = rep.get("file_contains", "")   # 空=全docx対象／指定=そのファイル名を含むもののみ
        desc = rep.get("desc", "")
        tok = TX.token_for(src)
        if tok:
            pairs.append(("{{%s}}" % tok, val, desc, fc))   # 主：トークン穴埋め
        finds = rep.get("find", [])
        if isinstance(finds, str):
            finds = [finds]
        for find in finds:                                  # 副：サンプル値フォールバック
            if find and val != find:
                pairs.append((find, val, desc, fc))
    pairs.sort(key=lambda t: len(t[0]), reverse=True)
    return pairs


def _xml_escape(s):
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _col_num(col):
    n = 0
    for ch in col:
        n = n * 26 + (ord(ch) - 64)
    return n


def _write_xlsx_cells_xmlsurgery(xlsx_path, sheet_title, coord_values, run_log):
    """openpyxl保存を使わず、セル値だけをシートXMLへ直接書き込む（図形・customXml等を保持）。
       既存 <c> は inlineStr で更新、無ければ列順に挿入。他パーツはバイトコピー。"""
    import re as _re
    import zipfile as _zip
    if not coord_values:
        return
    try:
        z = _zip.ZipFile(xlsx_path, "r")
        names = z.namelist()
        data = {n: z.read(n) for n in names}
        z.close()
    except Exception as e:
        run_log.append("[xmlsurgery] 読込失敗、openpyxl保存にフォールバック: %r" % e)
        raise

    wbxml = data.get("xl/workbook.xml", b"").decode("utf-8", "ignore")
    relsxml = data.get("xl/_rels/workbook.xml.rels", b"").decode("utf-8", "ignore")
    # シート名 → r:id → ターゲットXML
    sheet_path = None
    rid = None
    for tag in _re.findall(r"<sheet\b[^>]*/>", wbxml):
        nm = _re.search(r'name="([^"]*)"', tag)
        ri = _re.search(r'r:id="([^"]*)"', tag)
        if nm and ri and nm.group(1) == sheet_title:
            rid = ri.group(1); break
    if rid is None:
        m = _re.search(r'<sheet\b[^>]*r:id="([^"]*)"', wbxml)
        rid = m.group(1) if m else None
    if rid:
        rm = _re.search(r'<Relationship\b[^>]*Id="%s"[^>]*Target="([^"]*)"' % _re.escape(rid), relsxml)
        if rm:
            tgt = rm.group(1).lstrip("/")
            sheet_path = tgt if tgt.startswith("xl/") else "xl/" + tgt
    if not sheet_path or sheet_path not in data:
        cand = [n for n in names if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
        if not cand:
            raise RuntimeError("シートXMLが見つからない")
        sheet_path = sorted(cand)[0]

    xml = data[sheet_path].decode("utf-8")

    def _cell_xml(coord, s_attr, val):
        if val is None or val == "":
            return '<c r="%s"%s/>' % (coord, s_attr)
        return '<c r="%s"%s t="inlineStr"><is><t xml:space="preserve">%s</t></is></c>' % (
            coord, s_attr, _xml_escape(val))

    missing = {}
    for coord, val in coord_values.items():
        pat = _re.compile(r'<c r="%s"([^>/]*)(/>|>.*?</c>)' % _re.escape(coord), _re.S)
        m = pat.search(xml)
        if m:
            sm = _re.search(r'\ss="(\d+)"', m.group(1))
            s_attr = ' s="%s"' % sm.group(1) if sm else ""
            xml = xml[:m.start()] + _cell_xml(coord, s_attr, val) + xml[m.end():]
        else:
            missing[coord] = val

    # 既存行に無いセルは列順で挿入（無ければ行ごと挿入）
    for coord, val in missing.items():
        mm = _re.match(r"([A-Z]+)(\d+)", coord)
        if not mm:
            continue
        col, rownum = mm.group(1), int(mm.group(2))
        cellxml = _cell_xml(coord, "", val)
        rowpat = _re.compile(r'(<row r="%d"[^>]*>)(.*?)(</row>)' % rownum, _re.S)
        rm2 = rowpat.search(xml)
        if rm2:
            inner = rm2.group(2)
            # 列順の挿入位置を探す
            pos = None
            for cm in _re.finditer(r'<c r="([A-Z]+)\d+"', inner):
                if _col_num(cm.group(1)) > _col_num(col):
                    pos = cm.start(); break
            inner2 = (inner[:pos] + cellxml + inner[pos:]) if pos is not None else (inner + cellxml)
            xml = xml[:rm2.start()] + rm2.group(1) + inner2 + rm2.group(3) + xml[rm2.end():]
        else:
            # 行が無い → sheetData末尾へ（Excelは順不同を許容）
            newrow = '<row r="%d">%s</row>' % (rownum, cellxml)
            xml = xml.replace("</sheetData>", newrow + "</sheetData>", 1)

    data[sheet_path] = xml.encode("utf-8")

    # inlineStr化で共有文字列への参照数が減るため、sharedStrings の count を実数に合わせる
    # （count不整合＝Excelの「修復されたレコード」警告の原因）。sharedStrings本体は改変しない
    # （未参照のトークン文字列が残るが表示されない。空<t>化＝別の破損原因になるため触らない）。
    ss_key = "xl/sharedStrings.xml"
    if ss_key in data:
        total_refs = 0
        for n in names:
            if _re.match(r"xl/worksheets/sheet\d+\.xml$", n):
                content = xml if n == sheet_path else data[n].decode("utf-8", "ignore")
                total_refs += len(_re.findall(r'<c\b[^>]*\bt="s"', content))
        ss = data[ss_key].decode("utf-8")
        ss = _re.sub(r'(<sst\b[^>]*\bcount=")\d+(")', r"\g<1>%d\g<2>" % total_refs, ss, count=1)
        data[ss_key] = ss.encode("utf-8")

    tmp = xlsx_path + ".tmp"
    zo = _zip.ZipFile(tmp, "w", _zip.ZIP_DEFLATED)
    try:
        for n in names:
            zo.writestr(n, data[n])
    finally:
        zo.close()
    os.replace(tmp, xlsx_path)


def _resolve_sheet_xml_path(data, names, sheet_title):
    """workbook.xml/rels から sheet_title に対応するワークシートXMLのパスを返す。"""
    import re as _re
    wbxml = data.get("xl/workbook.xml", b"").decode("utf-8", "ignore")
    relsxml = data.get("xl/_rels/workbook.xml.rels", b"").decode("utf-8", "ignore")
    rid = None
    for tag in _re.findall(r"<sheet\b[^>]*/>", wbxml):
        nm = _re.search(r'name="([^"]*)"', tag)
        ri = _re.search(r'r:id="([^"]*)"', tag)
        if nm and ri and nm.group(1) == sheet_title:
            rid = ri.group(1); break
    if rid is None:
        m = _re.search(r'<sheet\b[^>]*r:id="([^"]*)"', wbxml)
        rid = m.group(1) if m else None
    if rid:
        rm = _re.search(r'<Relationship\b[^>]*Id="%s"[^>]*Target="([^"]*)"' % _re.escape(rid), relsxml)
        if rm:
            tgt = rm.group(1).lstrip("/")
            sp = tgt if tgt.startswith("xl/") else "xl/" + tgt
            if sp in data:
                return sp
    cand = [n for n in names if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
    return sorted(cand)[0] if cand else None


def _delete_rows_xmlsurgery(xlsx_path, sheet_title, start, count, run_log):
    """指定シートから連続行 [start, start+count-1] を XML手術で削除（図形・書式・データ検証を保持）。
       行番号・セル参照を繰り上げ、mergeCells と dimension を更新する。openpyxl は使わない。"""
    import re as _re
    import zipfile as _zip
    if count <= 0:
        return
    end = start + count - 1
    try:
        z = _zip.ZipFile(xlsx_path, "r")
        names = z.namelist()
        data = {n: z.read(n) for n in names}
        z.close()
    except Exception as e:
        run_log.append("[行削除] 読込失敗: %r" % e); return
    key = _resolve_sheet_xml_path(data, names, sheet_title)
    if not key:
        run_log.append("[行削除] シートXML未検出"); return
    xml = data[key].decode("utf-8")

    m = _re.search(r"(<sheetData>)(.*)(</sheetData>)", xml, _re.S)
    if not m:
        run_log.append("[行削除] sheetData未検出"); return
    head, body, tail = xml[:m.start()], m.group(2), xml[m.end():]
    out = []
    for el in _re.findall(r"<row\b[^>]*?(?:/>|>.*?</row>)", body, _re.S):
        mr = _re.search(r'<row\b[^>]*?\br="(\d+)"', el)
        if not mr:
            out.append(el); continue
        rn = int(mr.group(1))
        if start <= rn <= end:
            continue
        if rn > end:
            newn = rn - count
            el = _re.sub(r'(<row\b[^>]*?\br=")%d(")' % rn,
                         lambda mm: mm.group(1) + str(newn) + mm.group(2), el, count=1)
            el = _re.sub(r'(r=")([A-Z]+)%d(")' % rn,
                         lambda mm: mm.group(1) + mm.group(2) + str(newn) + mm.group(3), el)
        out.append(el)
    xml = head + "<sheetData>" + "".join(out) + "</sheetData>" + tail

    mm = _re.search(r'<mergeCells count="\d+">(.*?)</mergeCells>', xml, _re.S)
    if mm:
        def _rc(c):
            mo = _re.match(r"([A-Z]+)(\d+)", c); return mo.group(1), int(mo.group(2))
        newrefs = []
        for ref in _re.findall(r'<mergeCell ref="([^"]+)"/>', mm.group(1)):
            has = ":" in ref
            a, b = ref.split(":") if has else (ref, ref)
            ca, ra = _rc(a); cb, rb = _rc(b)
            if start <= ra <= end and start <= rb <= end:
                continue
            ra2 = ra - count if ra > end else ra
            rb2 = rb - count if rb > end else rb
            newrefs.append("%s%d:%s%d" % (ca, ra2, cb, rb2) if has else "%s%d" % (ca, ra2))
        xml = xml[:mm.start()] + '<mergeCells count="%d">%s</mergeCells>' % (
            len(newrefs), "".join('<mergeCell ref="%s"/>' % r for r in newrefs)) + xml[mm.end():]

    dm = _re.search(r'<dimension ref="([A-Z]+)(\d+):([A-Z]+)(\d+)"/>', xml)
    if dm:
        xml = xml[:dm.start()] + '<dimension ref="%s%s:%s%d"/>' % (
            dm.group(1), dm.group(2), dm.group(3), int(dm.group(4)) - count) + xml[dm.end():]

    # 手動改ページ（rowBreaks）を行削除に合わせて更新（範囲内は除去、下は -count）
    rbm = _re.search(r'(<rowBreaks\b[^>]*>)(.*?)(</rowBreaks>)', xml, _re.S)
    if rbm:
        newbrks = []
        for b in _re.findall(r'<brk\b[^>]*/>', rbm.group(2)):
            im = _re.search(r'\bid="(\d+)"', b)
            if not im:
                newbrks.append(b); continue
            rid = int(im.group(1))
            if start <= rid <= end:
                continue                                   # 削除範囲内の改ページは除去
            if rid > end:
                b = _re.sub(r'(\bid=")\d+(")', lambda mm: mm.group(1) + str(rid - count) + mm.group(2), b, count=1)
            newbrks.append(b)
        hdr = _re.sub(r'\bcount="\d+"', 'count="%d"' % len(newbrks), rbm.group(1))
        hdr = _re.sub(r'\bmanualBreakCount="\d+"', 'manualBreakCount="%d"' % len(newbrks), hdr)
        xml = xml[:rbm.start()] + hdr + "".join(newbrks) + rbm.group(3) + xml[rbm.end():]

    data[key] = xml.encode("utf-8")
    tmp = xlsx_path + ".tmp"
    zo = _zip.ZipFile(tmp, "w", _zip.ZIP_DEFLATED)
    try:
        for n in names:
            zo.writestr(n, data[n])
    finally:
        zo.close()
    os.replace(tmp, xlsx_path)
    run_log.append("[行削除] %s の 行%d〜%d (%d行) を削除" % (sheet_title, start, end, count))


def _docx_items(docx_path):
    """説明書docx（4.5等）を「N. 見出し」で項目分割し {N: 本文テキスト} を返す。
       本文は見出し行の次から次の見出し手前までの段落を改行連結。"""
    import re as _re
    try:
        from docx import Document
    except Exception:
        return {}
    if not docx_path or not os.path.exists(docx_path):
        return {}
    try:
        paras = [_clean(p.text) for p in Document(docx_path).paragraphs]
    except Exception:
        return {}
    heads = {}
    for i, t in enumerate(paras):
        m = _re.match(r"^\s*(\d+)[\.．]\s*(.+)", (t or "").strip())
        if m and len(m.group(2)) < 40:
            try:
                heads[int(m.group(1))] = i
            except ValueError:
                pass
    order = sorted(heads)
    items = {}
    for k in order:
        i = heads[k]
        nxt = min([heads[j] for j in order if heads[j] > i] + [len(paras)])
        items[k] = "\n".join(t for t in paras[i + 1:nxt] if t and t.strip())
    return items


def _fill_xlsx_tokens(ws, hearing, kits, TX, saisei, doc_items=None):
    """様式xlsxセル内に残る {{項目名}} をヒアリング値で穴埋め（部分セル・旧トークン対応）。
       特殊トークン（採血量/治療価格/キット製造方法/委員会/場所/4.5本文）を名前で解決し、
       それ以外は「より→から」等を吸収してヒアリングlabel直接参照でフォールバック。
       doc_items: 説明書docxの項目辞書 {N: 本文}（{{4.5本文:N}} 用）。"""
    import re as _re
    TOKENRE = _re.compile(r"\{\{([^{}]*)\}\}")
    MARKS = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮"

    def _kit_index(name):
        m = _re.search(r"メーカー([①-⑮])", name)
        if m:
            return MARKS.find(m.group(1))
        m = _re.search(r"メーカー([A-O])", name)
        if m:
            return ord(m.group(1)) - ord("A")
        return 0

    def _resolve_named(name):
        name = name.strip()
        n2 = name.replace("より", "から")           # 旧表記ゆれ吸収
        # 4.5等 説明書の項目本文（{{4.5本文:1}} = 説明書の項目1本文）を差込
        # ※番号は「本文」の後を取る（「4.5」の4を拾わないように）
        if name.startswith("4.5本文") or name.startswith("4.5 本文") or "説明書本文" in name:
            m = _re.search(r"本文[^\d]*(\d+)", name)
            if m and doc_items:
                return doc_items.get(int(m.group(1)), "")
            return ""
        if name == "必要な採血量":
            return hearing.blood_volume()
        if name == "治療価格":
            return hearing.kit_price_block(saisei)
        if "治療費の設定" in name:
            # 様式は「①キット名を用いた治療：{価格}」形式のため、ブロックではなく単一価格を返す
            idx = 2 if "②" in name else (3 if "③" in name else 1)
            plabel = "再生医療%s（右記タブから選択）の治療費の設定（税別）" % MARKS[idx - 1]
            return TX.clean(hearing.lookup(plabel, "", 1)) or ""
        # キット製造方法：スロット番号（①/A等）付きは「その1キット単体」、無印は全キットのブロック
        def _kit_method(part):
            ks = hearing.prp_kits()
            i = _kit_index(name)
            has_mark = bool(_re.search(r"メーカー([①-⑮A-O])", name))
            if has_mark:
                return hearing.kit_method(ks[i], part) if i < len(ks) else ""
            return hearing.kit_block(part, numbered=True)
        if name == "採取方法" or ("採取" in name and "キット" in name):
            return _kit_method("採取")
        if name == "加工方法" or ("加工の方法" in name):
            return _kit_method("加工")
        if name == "投与方法" or ("投与の方法" in name):
            return _kit_method("投与")
        if "PRPキットメーカー" in name and "右記タブ" in name:
            ks = hearing.prp_kits(); i = _kit_index(name)
            return _re.sub(r"\s*（[^（）]*）\s*$", "", ks[i]).strip() if i < len(ks) else ""
        if "認定再生医療等委員会" in name and "認定番号" in name:
            cname = hearing.lookup("認定再生医療等委員会の名称", "審査委員会", 1)
            return hearing.committee_field(cname, 2)
        # セクション:ラベル 形式（連絡先の氏名/電話番号/メールアドレス等）
        sv = TX.section_label_value(name, hearing)
        if sv:
            return sv
        # 汎用：ヒアリングlabel直接参照（採血場所/投与場所/保険名称/問い合わせ先/委員会名称/電話/メール等）
        v = hearing.lookup(n2, "", 1)
        return TX.clean(v) if v not in (None, "") else ""

    n = 0
    for row in ws.iter_rows():
        for c in row:
            if not isinstance(c.value, str) or "{{" not in c.value:
                continue
            def _rep(m):
                val = _resolve_named(m.group(1))
                return val if val else m.group(0)   # 未解決は残す（後段のclear_tokensで空欄化）
            newv = TOKENRE.sub(_rep, c.value)
            # 採血量の注記がテンプレ固定文と値で重複した場合は1つに畳む
            newv = _re.sub(r"(（使用キットにより異なる）)[（(]使用キットにより異なる[）)]", r"\1", newv)
            if newv != c.value:
                TX.safe_set(ws, c.coordinate, newv)
                n += 1
    return n


def _doctor_token_value(hearing, token_name):
    """docx の医師トークン（例: {{再生医療を行う医師の役職・氏名; 氏名1}}）を
       「接尾の番号 N の医師」1人分に解決する。氏名N→N人目の氏名 / 役職N→N人目の役職。
       番号なしは1人目。ヒアリング値は「役職　氏名」（全角空白区切り）。該当なしは ''。"""
    import re as _re
    suffix = _re.split(r"[;:；：]", token_name)[-1] if _re.search(r"[;:；：]", token_name) else token_name
    mnum = _re.search(r"(\d+)", suffix)
    idx = int(mnum.group(1)) if mnum else 1
    want_role = ("役職" in suffix) and ("氏名" not in suffix)
    docs = hearing.doctors()
    if idx < 1 or idx > len(docs):
        return ""
    d = (docs[idx - 1] or "").strip()
    role, name = "", d
    if "　" in d:
        role, name = [s.strip() for s in d.split("　", 1)]
    return role if want_role else name


def _fill_paragraph_multiline_tokens(p, resolved, run_green):
    """段落内の {{名前}} を resolved[名前] で置換し、値内の改行を <w:br/> で保持したまま
       段落を作り直す。接頭辞（①等）や前後の定型文は残し、差込値のみ緑にする。
       文中インラインだが値が複数行のケース（例: ①{{…細胞採取の方法}}）で改行を守る。"""
    import re as _re
    TOKEN = _re.compile(r"\{\{([^{}]*)\}\}")
    full = "".join(r.text for r in p.runs)
    segs = []          # ("text", 文字列, 緑か) / ("br",)
    last = 0
    replaced = False
    for m in TOKEN.finditer(full):
        nm = m.group(1).strip()
        if nm not in resolved:
            continue
        if m.start() > last:                       # トークン前の定型テキスト
            segs.append(("text", full[last:m.start()], False))
        lines = resolved[nm].split("\n")
        segs.append(("text", lines[0], True))
        for ln in lines[1:]:                        # 2行目以降は改行(br)で継ぐ
            segs.append(("br",))
            segs.append(("text", ln, True))
        last = m.end()
        replaced = True
    if not replaced:
        return 0
    if last < len(full):                            # トークン後の定型テキスト
        segs.append(("text", full[last:], False))
    for r in list(p.runs):                          # 既存ランを除去して作り直す
        r._element.getparent().remove(r._element)
    run = None
    for seg in segs:
        if seg[0] == "br":
            if run is None:
                run = p.add_run("")
            run.add_break()
        else:
            run = p.add_run(seg[1])
            if seg[2]:
                run_green(run)
    return 1


def _fill_all_docx_tokens(d, hearing, kits, saisei, TX, iter_paras):
    """docx内に残る任意の {{項目名}} を汎用リゾルバで穴埋め（前世代トークンの取りこぼし防止）。
       段落全体が単一トークン→複数行(改行br)で差込／文中インライン→ラン置換（1行化）。"""
    import re as _re
    from transcribe_sop import _run_green, _replace_in_paragraph
    TOKEN = _re.compile(r"\{\{([^{}]*)\}\}")
    STANDALONE = _re.compile(r"^\s*\{\{([^{}]*)\}\}\s*$")
    cache = {}

    def val_for(nm):
        if nm not in cache:
            # 医師の役職・氏名トークンは接尾の番号N の医師1人に解決（xlsxは doctor_fill が個別処理）
            if "再生医療を行う医師" in nm and ("氏名" in nm or "役職" in nm):
                cache[nm] = _doctor_token_value(hearing, nm)
            else:
                cache[nm] = TX.resolve_token_by_name(nm, hearing, kits, saisei) or ""
        return cache[nm]

    n = 0
    for p in iter_paras(d):
        if "{{" not in p.text:
            continue
        m = STANDALONE.match(p.text)
        if m:
            v = val_for(m.group(1).strip())
            if v == "":
                continue                       # 空は _clear_doc_tokens に任せる
            lines = v.split("\n")
            # 元ラン（テンプレの{{…}}を保持していたラン）に明示サイズがあれば控えておく。
            # add_run で作る新ランはサイズ無指定＝文書既定(10.5pt)に戻ってしまうため、
            # 明示サイズを新ランへ引き継いで見た目を保つ（例: 日付トークンの12pt維持）。
            src_sz = _run_sz_halfpts(p.runs[0]._element, default=None) if p.runs else None
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(lines[0]); _run_green(run); _apply_run_sz(run, src_sz)
            for ln in lines[1:]:
                run.add_break(); run = p.add_run(ln); _run_green(run); _apply_run_sz(run, src_sz)
            n += 1
        else:
            resolved = {}
            for nm in set(x.strip() for x in TOKEN.findall(p.text)):
                v = val_for(nm)
                if v != "":
                    resolved[nm] = v
            if not resolved:
                continue
            if any("\n" in v for v in resolved.values()):
                # 値が複数行（例: ①{{…細胞採取の方法}}）→ 改行を<w:br/>で保持して再構築
                n += _fill_paragraph_multiline_tokens(p, resolved, _run_green)
            else:
                pairs = [("{{%s}}" % nm, v) for nm, v in resolved.items()]
                n += _replace_in_paragraph(p, pairs)
    return n


def _run_sz_halfpts(r_el, default=36):
    """ラン要素の rPr から文字サイズ（ハーフポイント）を取得。無ければ default。"""
    from docx.oxml.ns import qn
    rpr = r_el.find(qn("w:rPr"))
    if rpr is not None:
        sz = rpr.find(qn("w:sz"))
        if sz is not None and sz.get(qn("w:val")):
            try:
                return int(sz.get(qn("w:val")))
            except ValueError:
                pass
    return default


def _apply_run_sz(run, half_pts):
    """ハーフポイント値をランへ設定（w:sz と w:szCs の両方）。None/0 のときは何もしない。"""
    if not half_pts:
        return
    try:
        from docx.shared import Pt
        run.font.size = Pt(half_pts / 2.0)
    except Exception:
        pass


def _build_ruby_run(base, reading, sz, color):
    """ルビ入りラン <w:r><w:ruby>…</w:ruby></w:r> を生成して返す（本文=base／ルビ=reading）。
       sz は本文サイズ（ハーフポイント）。ルビは約半分。色は本文・ルビ共通で付与。"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    ruby_sz = max(10, int(round(sz * 24.0 / 52.0)))     # 既存ルビ(base52/ルビ24)の比率に合わせる
    raise_v = max(1, int(round(sz * 50.0 / 52.0)))
    col = '<w:color w:val="%s"/>' % color if color else ""
    b = _xml_escape(base)
    r = _xml_escape(reading)
    xml = (
        '<w:r %s>' % nsdecls("w") +
        '<w:rPr>%s<w:sz w:val="%d"/><w:szCs w:val="%d"/></w:rPr>' % (col, sz, sz) +
        '<w:ruby>'
        '<w:rubyPr>'
        '<w:rubyAlign w:val="distributeSpace"/>'
        '<w:hps w:val="%d"/>' % ruby_sz +
        '<w:hpsRaise w:val="%d"/>' % raise_v +
        '<w:hpsBaseText w:val="%d"/>' % sz +
        '<w:lid w:val="ja-JP"/>'
        '</w:rubyPr>'
        '<w:rt><w:r><w:rPr>%s<w:sz w:val="%d"/><w:szCs w:val="%d"/></w:rPr>'
        '<w:t xml:space="preserve">%s</w:t></w:r></w:rt>' % (col, ruby_sz, ruby_sz, r) +
        '<w:rubyBase><w:r><w:rPr>%s<w:sz w:val="%d"/><w:szCs w:val="%d"/></w:rPr>'
        '<w:t xml:space="preserve">%s</w:t></w:r></w:rubyBase>' % (col, sz, sz, b) +
        '</w:ruby>'
        '</w:r>'
    )
    return parse_xml(xml)


def _fill_token_ruby(d, iter_paras, token, base, reading, color="00B050"):
    """token を含む段落を、base（本文）に reading（ふりがな）を付けたルビ入りランへ置換する。
       元ランの文字サイズを引き継ぐ。1件でも置換したら件数を返す（0=対象なし）。"""
    n = 0
    for p in iter_paras(d):
        if token not in p.text:
            continue
        sz = _run_sz_halfpts(p.runs[0]._element, 36) if p.runs else 36
        for r in list(p.runs):                       # トークンの既存ランを除去
            r._element.getparent().remove(r._element)
        p._p.append(_build_ruby_run(base, reading, sz, color))
        n += 1
    return n


def _clear_doc_tokens(d, iter_paras):
    """穴埋めされずに残った {{...}} を空欄化（未充足トークンをクライアントに見せない）。"""
    import re as _re
    TOKEN_RE = _re.compile(r"\{\{[^{}]*\}\}")
    n = 0
    for p in iter_paras(d):
        for r in p.runs:
            if r.text and "{{" in r.text and "}}" in r.text:
                nt = TOKEN_RE.sub("", r.text)
                if nt != r.text:
                    r.text = nt
                    n += 1
        # ラン跨ぎの残存トークンも段落テキストで検出して除去
        full = "".join(r.text for r in p.runs)
        if "{{" in full and "}}" in full and p.runs:
            new = TOKEN_RE.sub("", full)
            if new != full:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
                n += 1
    return n


def _rireki_sensei(path):
    """入力略歴書のファイル名から先生名を取り出す（例 略歴書_大井知泉先生.xlsx → 大井知泉先生）。"""
    import re as _re
    b = os.path.splitext(os.path.basename(path))[0]
    m = _re.search(r"略歴書[_＿](.+)$", b)
    return m.group(1) if m else b


def _fill_one_rireki(inp_path, out_path):
    """入力略歴書xlsx（A列ラベル/B列値）から出力略歴書xlsx を埋める。
       学歴/医師免許/職歴 は入力側で混在しうるため内容判定で振り分ける。図形なし前提でopenpyxl保存。"""
    import openpyxl
    inws = openpyxl.load_workbook(inp_path, data_only=True).worksheets[0]
    owb = openpyxl.load_workbook(out_path)
    ows = owb.worksheets[0]
    rows = [(_clean(inws.cell(r, 1).value), inws.cell(r, 2).value) for r in range(1, inws.max_row + 1)]
    allb = [b for a, b in rows if b not in (None, "")]

    def by_label(*keys):
        for a, b in rows:
            if a and any(k in a for k in keys) and b not in (None, ""):
                return b
        return None

    def containing(*keys):
        for b in allb:
            if any(k in _clean(b) for k in keys):
                return b
        return None

    menkyo = containing("医籍番号", "免許取得日")
    gs = by_label("医師免許", "職歴")        # 学歴+職歴が混在しうる入力ブロック
    gakureki = shokureki = None
    if gs:
        lines = [l for l in _clean(gs).split("\n") if l.strip()]
        g = [l for l in lines if ("大学" in l and "卒" in l)]
        s = [l for l in lines if not ("大学" in l and "卒" in l)]
        gakureki = "\n".join(g) if g else None
        shokureki = "\n".join(s) if s else None
    furi = name = None
    for i, (a, b) in enumerate(rows):
        if a and "氏名" in a:
            furi = b
            if i + 1 < len(rows):
                name = rows[i + 1][1]
            break

    def setc(row, val):
        if val not in (None, ""):
            ows.cell(row, 2).value = val
            return 1
        return 0

    filled = 0
    tr = 1
    while tr <= ows.max_row:
        lab = _clean(ows.cell(tr, 1).value)
        if lab.startswith("氏名"):
            filled += setc(tr, furi); filled += setc(tr + 1, name)
        elif "生年月日" in lab:
            filled += setc(tr, by_label("生年月日"))
        elif lab.startswith("所属") and "学会" not in lab:
            filled += setc(tr, by_label("所属"))
        elif "役職" in lab:
            filled += setc(tr, by_label("役職"))
        elif "学歴" in lab:
            filled += setc(tr, gakureki)
        elif "医師免許" in lab:
            filled += setc(tr, menkyo)
        elif lab == "職歴":
            filled += setc(tr, shokureki)
        elif "専門分野" in lab:
            filled += setc(tr, by_label("専門分野"))
        elif "所属学会" in lab:
            filled += setc(tr, by_label("所属学会"))
        elif "認定医" in lab or "資格" in lab:
            filled += setc(tr, by_label("認定医", "資格"))
        elif "臨床経験" in lab:
            idx = None
            for i, (a, b) in enumerate(rows):
                if a and "臨床経験" in a:
                    idx = i; break
            if idx is not None:
                vals = [b for (a, b) in rows[idx:] if b not in (None, "")]
                for k, v in enumerate(vals):
                    filled += setc(tr + k, v)
        tr += 1
    owb.save(out_path)
    return filled


def _generate_rireki_outputs(out_folder, input_dir, run_log):
    """出力フォルダの略歴書テンプレ（3.医師略歴書*.xlsx）を、01_inputの各略歴書ごとに埋めて出力。
       入力ファイル数ぶんの 3.医師略歴書_<先生名>.xlsx を生成し、未充填のテンプレ複製は削除。"""
    import shutil as _sh
    tpls = [f for f in glob.glob(os.path.join(out_folder, "*.xlsx"))
            if os.path.basename(f).startswith("3.医師略歴書")]
    if not tpls:
        return
    tpl = tpls[0]
    inputs = [f for f in glob.glob(os.path.join(input_dir, "*.xlsx"))
              if "略歴" in os.path.basename(f) and not os.path.basename(f).startswith("~$")]
    if not inputs:
        return
    made = []
    for inp in sorted(inputs):
        sensei = _rireki_sensei(inp)
        out = os.path.join(out_folder, "3.医師略歴書_%s.xlsx" % sensei)
        try:
            _sh.copyfile(tpl, out)
            _fill_one_rireki(inp, out)
            made.append(sensei)
        except Exception as e:
            run_log.append("[略歴書] %s の生成失敗: %r" % (sensei, e))
    # テンプレ複製（自分が作った出力名でなければ）を削除
    if os.path.basename(tpl) not in ["3.医師略歴書_%s.xlsx" % m for m in made] and os.path.exists(tpl):
        try:
            os.remove(tpl)
        except Exception:
            pass
    run_log.append("[略歴書] 医師%d名分を生成: %s" % (len(made), " / ".join(made)))


def run_docs(hearing, hearing_path, dir_tpl, dir_output, run_log, run_dt):
    """フォルダ型テンプレ一式を転記出力。戻り値: (出力フォルダ一覧, log_rows)。"""
    try:
        from docx import Document
    except ImportError:
        run_log.append("[docs] python-docx 未導入のためWord書類はスキップ")
        Document = None
    import openpyxl
    import transcribe as TX
    try:
        from transcribe_sop import (_replace_in_paragraph, _iter_all_paragraphs,
                                     _replace_lowlevel, _fill_token_multiline, _set_cell_multiline,
                                     reset_doc_green_to_black)
    except Exception as e:
        run_log.append("[docs] Word置換関数の読込失敗: %r" % e)
        _replace_in_paragraph = _iter_all_paragraphs = _replace_lowlevel = None
        _fill_token_multiline = _set_cell_multiline = None

    if not os.path.exists(DOCS_CONFIG):
        run_log.append("[docs] docs_config.json が無いためスキップ")
        return [], []
    cfg = json.load(io.open(DOCS_CONFIG, encoding="utf-8"))

    kits = hearing.prp_kits()
    furi = _load_furigana()          # 医療機関名→ルビ（手入力マップ）
    base = os.path.splitext(os.path.basename(hearing_path))[0]
    out_folders = []
    log_rows = []
    green_report = []   # 元が緑ターゲットだが未転記だったセル（要確認候補）

    for doc_key, doc in cfg.get("documents", {}).items():
        tpl_folder = os.path.join(dir_tpl, doc["folder"])
        if not os.path.isdir(tpl_folder):
            run_log.append("[%s] テンプレフォルダ無し: %s" % (doc_key, tpl_folder))
            continue
        out_folder = os.path.join(dir_output, "%s_%s" % (doc["folder"], base))
        try:
            if os.path.exists(out_folder):
                shutil.rmtree(_long(out_folder))
            # 拡張パス（\\?\）で複製し、深い階層の長いファイル名でも 260 文字上限に阻まれないようにする
            shutil.copytree(_long(tpl_folder), _long(out_folder))   # PDF・論文・追加xlsx等も無変更で同梱
        except PermissionError as e:
            run_log.append("[%s] ★スキップ：出力先のファイルが開かれています。"
                           "02_output内のWord/Excel（特に前回の出力）を閉じてから再実行してください。（%s）"
                           % (doc_key, e))
            continue
        except Exception as e:
            run_log.append("[%s] コピー失敗でスキップ: %r" % (doc_key, e))
            continue

        # --- 様式Excel（セル転記） ---
        n_x_done = n_x_check = 0
        exc = doc.get("excel")
        if exc:
            xpath = os.path.join(out_folder, exc["file"])
            if os.path.exists(xpath):
                wb = openpyxl.load_workbook(xpath)
                sheet = exc.get("sheet")
                ws = wb[sheet] if sheet in wb.sheetnames else wb[wb.sheetnames[0]]
                # 図形保持のため、書込みは最後にXML手術で行う。転記前の全セル値を控える。
                orig_vals = {c.coordinate: c.value for r in ws.iter_rows() for c in r if c.value is not None}
                # 転記前：テンプレの緑セルを記録し、一旦すべて黒に戻す
                orig_green = {}
                for row in ws.iter_rows():
                    for c in row:
                        if _is_green_cell(c):
                            orig_green[c.coordinate] = _clean(c.value)
                            _set_cell_color(c, BLACK)
                written = set()   # 転記したセル（後で緑にする）
                for e in exc.get("entries", []):
                    if e.get("_disabled"):
                        continue
                    val, st = TX.resolve(e["source"], hearing, kits)
                    if st == TX.ST_UNMAP or val in (None, ""):
                        continue
                    val = TX.apply_tf(val, e.get("tf", "text"))
                    if val in (None, ""):
                        continue
                    if TX.safe_set(ws, e["cell"], val):
                        written.add(e["cell"])
                        if st == TX.ST_CHECK or "確認対象" in _clean(e.get("note")):
                            n_x_check += 1
                        else:
                            n_x_done += 1
                        log_rows.append(TX.make_row(run_dt, doc_key, e, val, st, e.get("note")))
                # 医師の可変転記（最大N名。B列ラベルで医師ブロックを動的検出し、氏名/所属を順に）
                doctor_del = None      # 空の医師ブロック（末尾）の削除範囲 (start_row, count)
                df = exc.get("doctor_fill")
                if df:
                    from openpyxl.utils import get_column_letter as _gl
                    blabel = df.get("block_label", "")
                    col = int(df.get("col", 12))
                    noff = int(df.get("name_offset", 1))
                    aoff = df.get("affil_offset")
                    maxn = int(df.get("max", 15))
                    blocks = [r for r in range(1, ws.max_row + 1)
                              if isinstance(ws.cell(r, 2).value, str) and blabel and blabel in ws.cell(r, 2).value]
                    docs = hearing.doctors()
                    med, _mst = TX.resolve({"t": "hearing",
                                            "label": "医療機関/名称（診療所開設届上）",
                                            "section": "法人/医療機関"}, hearing, kits)
                    roff = df.get("role_offset")     # 役職欄のオフセット（あれば役職・氏名を分離）
                    for i, dname in enumerate(docs[:maxn]):
                        if i >= len(blocks):
                            run_log.append("[%s] 医師%d名目以降はテンプレの医師ブロック不足で未転記（ブロック追加が必要）"
                                           % (doc_key, i + 1))
                            break
                        br = blocks[i]
                        # ヒアリングは「役職　氏名」（全角空白区切り）。転記時に分離する。
                        role, name = "", _clean(dname)
                        if roff is not None and "　" in name:
                            role, name = [s.strip() for s in name.split("　", 1)]
                        nc = "%s%d" % (_gl(col), br + noff)
                        TX.safe_set(ws, nc, name); written.add(nc)
                        if aoff is not None and med:
                            ac = "%s%d" % (_gl(col), br + int(aoff))
                            TX.safe_set(ws, ac, med); written.add(ac)
                        if roff is not None and role:
                            rc = "%s%d" % (_gl(col), br + int(roff))
                            TX.safe_set(ws, rc, role); written.add(rc)
                        n_x_done += 1
                    if docs:
                        run_log.append("[%s] 医師転記=%d名（テンプレ医師ブロック=%d）"
                                       % (doc_key, min(len(docs), len(blocks)), len(blocks)))
                    # 空の医師ブロック（末尾）の行を出力から削除する指定（delete_empty:true）
                    if df.get("delete_empty") and blocks:
                        nfilled = min(len(docs), len(blocks))
                        if nfilled < len(blocks):
                            brows = (blocks[1] - blocks[0]) if len(blocks) > 1 else 4
                            doctor_del = (blocks[nfilled], (len(blocks) - nfilled) * brows)
                # チェックボックス選択（ラジオ）：ヒアリング値に一致する選択肢の□を■に
                #   例) 救急 自施設/他施設 … ヒアリング「他の医療機関」→ 他の医療機関側を■
                for cbs in exc.get("checkbox_selects", []):
                    cval, _cst = TX.resolve(cbs.get("source", {}), hearing, kits)
                    cval = _clean(cval)
                    on = cbs.get("on", "■")
                    off = cbs.get("off", "□")
                    picked = None
                    for opt in cbs.get("options", []):
                        m = _clean(opt.get("match", ""))
                        hit = bool(cval) and (m in cval or cval in m)
                        mark = on if hit else off
                        if TX.safe_set(ws, opt["cell"], mark):
                            written.add(opt["cell"])
                        if hit:
                            picked = m
                    run_log.append("[%s] %s → %s" % (doc_key, cbs.get("desc", "選択"),
                                                     ("%s を■" % picked) if picked else "該当なし(全て□)"))
                # セル内の一部だけ置換（保険名称/問い合わせ先/採取方法など。定型文は保持）
                for ce in exc.get("cell_edits", []):
                    cur = ws[ce["cell"]].value
                    cur = "" if cur is None else str(cur)
                    val, _st = TX.resolve(ce.get("source", {}), hearing, kits)
                    val = _clean(val)
                    if not val:
                        continue
                    op = ce.get("op")
                    new = cur
                    # トークン方式：セル内の {{項目名}} を先に穴埋め（部分セルのトークン化対応）
                    _tok = TX.token_for(ce.get("source", {}))
                    if _tok and ("{{%s}}" % _tok) in new:
                        new = new.replace("{{%s}}" % _tok, val)
                    if op == "replace_find":            # 指定文字列を丸ごと置換
                        for fnd in ce.get("find", []):
                            if fnd and fnd in new:
                                new = new.replace(fnd, val)
                    elif op == "replace_from":          # マーカー以降を置換（前は保持）
                        marker = ce.get("marker", "")
                        idx = cur.find(marker) if marker else -1
                        if idx >= 0:
                            new = cur[:idx] + val
                    elif op == "replace_price":         # 治療費の①②…行を「キット×治療費」ブロックへ置換
                        lines = new.split("\n")
                        out = []; done = False
                        for ln in lines:
                            if "を用いた治療" in ln:      # 既存の価格行（サンプル/トークン問わず）
                                if not done:
                                    out.append(val); done = True
                            else:
                                out.append(ln)
                        new = "\n".join(out) if done else new
                    if new != cur and TX.safe_set(ws, ce["cell"], new):
                        written.add(ce["cell"])
                        n_x_done += 1
                        log_rows.append(TX.make_row(run_dt, doc_key,
                                                    {"sheet": exc.get("sheet"), "cell": ce["cell"],
                                                     "var": ce.get("desc", "cell_edit")},
                                                    val, TX.ST_CHECK, ce.get("desc")))
                # 転記したセルを緑にマーク（未転記判定はトークン穴埋め後に行う）
                for coord in written:
                    try:
                        _set_cell_color(ws[coord], GREEN)
                    except Exception:
                        pass
                # 部分セル・旧トークンの {{…}} をヒアリング値で穴埋め（clear_tokensの前に）
                _saisei = 2 if "3種" in doc_key else 1
                # {{4.5本文:N}} 用に説明書docx（4.5…）の項目本文を抽出（テンプレ側を参照＝最新内容）
                _exp = glob.glob(os.path.join(tpl_folder, "4.5*.docx")) or \
                       glob.glob(os.path.join(out_folder, "4.5*.docx"))
                _doc_items = _docx_items(_exp[0]) if _exp else {}
                _fill_xlsx_tokens(ws, hearing, kits, TX, _saisei, _doc_items)
                TX.clear_tokens(wb)   # なお残る未充足 {{…}} を空欄化（例: 3種様式の{{医師2_氏名}}）
                # 緑ターゲットの未転記判定は、トークン穴埋め・clear_tokens 後の「最終値」で行う。
                #   旧実装は穴埋め前に判定していたため、トークン方式で転記済みのセルまで
                #   「緑なのに未転記」に誤計上していた（＝本リストの精度低下の原因）。
                #   最終値が空、または未充足トークンが残るものだけを真の未転記として報告する。
                #   さらに、空の医師欄（delete_empty で行ごと削除される範囲）は意図的な空欄なので除外。
                del_lo = del_hi = None
                if doctor_del:
                    del_lo, del_hi = doctor_del[0], doctor_del[0] + doctor_del[1] - 1
                for coord, ov in orig_green.items():
                    if coord in written:
                        continue                     # entries/医師/チェックボックス等で転記済み
                    if del_lo is not None:
                        digits = "".join(ch for ch in coord if ch.isdigit())
                        if digits and del_lo <= int(digits) <= del_hi:
                            continue                 # 削除される空の医師欄行 → 報告しない
                    final = ws[coord].value
                    if final is None or str(final).strip() == "" or "{{" in str(final):
                        green_report.append((doc_key, sheet, coord, ov))
                    else:
                        try:                         # トークン等で値が入った緑ターゲットも緑に戻す
                            _set_cell_color(ws[coord], GREEN)
                        except Exception:
                            pass
                # 変更セルの差分を算出し、XML手術で書込み（openpyxl保存を避け図形を保持）
                diff = {}
                for r in ws.iter_rows():
                    for c in r:
                        if orig_vals.get(c.coordinate) != c.value:
                            diff[c.coordinate] = c.value
                try:
                    _write_xlsx_cells_xmlsurgery(xpath, ws.title, diff, run_log)
                except Exception as e:
                    run_log.append("[%s] XML手術失敗→openpyxl保存にフォールバック（図形が失われる可能性）: %r"
                                   % (doc_key, e))
                    wb.save(xpath)
                # 空の医師ブロックの行を出力から削除（値の書込み後・図形保持のXML手術）
                if doctor_del:
                    _delete_rows_xmlsurgery(xpath, ws.title, doctor_del[0], doctor_del[1], run_log)
            else:
                run_log.append("[%s] 様式xlsx無し: %s" % (doc_key, exc["file"]))

        # --- 直下の各docx（一括置換） ---
        n_files = n_repl = n_kit = n_ruby = 0
        if Document and _replace_in_paragraph:
            pairs = _build_docx_pairs(doc.get("docx", []), hearing, kits, TX)
            # キット製造方法トークン（{{採取方法}}等 → ①②③連番の複数行本文）
            kit_tokens = []
            for kt in doc.get("docx_kit", []):
                val, _ = TX.resolve(kt.get("source", {}), hearing, kits)
                kit_tokens.append((kt["token"], val or ""))
            # ルビ指定（医療機関名などにふりがなを付ける）。読みは手入力マップから引く。
            ruby_specs = []
            for rb in doc.get("docx_ruby", []):
                bval, _ = TX.resolve(rb.get("source", {}), hearing, kits)
                bval = _clean(bval)
                reading = _clean(furi.get(bval, ""))
                if bval and reading:                    # 読み未登録ならルビ無し（通常の穴埋めに任せる）
                    ruby_specs.append((rb["token"], bval, reading, rb.get("file_contains", "")))
                elif bval:
                    run_log.append("[%s] ルビ未登録：「%s」の読みが ふりがな.json に無いためルビ無しで出力"
                                   % (doc_key, bval))
            for f in sorted(glob.glob(os.path.join(out_folder, "*.docx"))):
                fbase = os.path.basename(f)
                if fbase.startswith("~$"):
                    continue
                try:
                    d = Document(f)
                except Exception as ex:
                    run_log.append("[%s] docx読込失敗 %s: %r" % (doc_key, fbase, ex))
                    continue
                reset_doc_green_to_black(d)   # テンプレの緑を一旦黒へ（転記箇所だけ後で緑になる）
                # ルビ付け（通常の穴埋めより先に。トークンをルビ入りランへ置換して以降の置換対象から外す）
                rc = 0
                for rtok, rbase, rreading, rfc in ruby_specs:
                    if rfc and rfc not in fbase:
                        continue
                    rc += _fill_token_ruby(d, _iter_all_paragraphs, rtok, rbase, rreading)
                n_ruby += rc
                # file_contains 指定があるものは対象ファイルのみに適用
                simple_pairs = [(fnd, v) for fnd, v, _, fc in pairs if (not fc or fc in fbase)]
                cnt = kc = 0
                for p in _iter_all_paragraphs(d):
                    cnt += _replace_in_paragraph(p, simple_pairs)
                cnt += _replace_lowlevel(d, simple_pairs)   # テキストボックス等も置換
                for tok, val in kit_tokens:                 # キットトークン差し込み
                    kc += _fill_token_multiline(d, tok, val)
                # 表フィル（略歴書など：行位置ベースでヒアリング略歴書シートから転記）
                for spec in doc.get("docx_table", []):
                    if spec.get("file_contains", "") not in fbase:
                        continue
                    ti = spec.get("table", 0); vcol = spec.get("value_col", 1)
                    tables = d.tables
                    if ti >= len(tables):
                        continue
                    tbl = tables[ti]
                    for rowdef in spec.get("rows", []):
                        ri = rowdef.get("row")
                        if ri is None or ri >= len(tbl.rows):
                            continue
                        val2, st2 = TX.resolve(rowdef.get("source", {}), hearing, kits)
                        if val2 in (None, ""):
                            continue
                        cells = tbl.rows[ri].cells
                        if vcol < len(cells):
                            try:
                                _set_cell_multiline(cells[vcol], val2)
                                kc += 1
                            except Exception as ex:
                                run_log.append("[docs] %s 表フィル失敗 r%s: %r" % (fbase, ri, ex))
                # 前世代を含む任意の {{項目名}} を汎用リゾルバで穴埋め（取りこぼし防止）
                _saisei_d = 2 if "3種" in doc_key else 1
                cnt += _fill_all_docx_tokens(d, hearing, kits, _saisei_d, TX, _iter_all_paragraphs)
                # 穴埋め後、なお残る未充足 {{...}} を空欄化（トークン方式の後始末）
                leftover = _clear_doc_tokens(d, _iter_all_paragraphs)
                if cnt or kc or leftover or rc:
                    d.save(f)
                    n_files += 1
                    n_repl += cnt
                    n_kit += kc
            # 置換サマリをログ行に
            for fnd, v, desc, fc in pairs:
                log_rows.append(TX.make_row(run_dt, doc_key,
                                            {"sheet": "(docx一括置換)", "cell": fnd, "var": desc},
                                            v, TX.ST_DONE, "サンプル値→ヒアリング値"))

        # 医師略歴書：01_input の各略歴書xlsxから 3.医師略歴書_<先生名>.xlsx を医師ごとに生成
        _generate_rireki_outputs(out_folder, os.path.dirname(hearing_path), run_log)

        out_folders.append(out_folder)
        run_log.append("[%s] 出力フォルダ: %s ／ 様式転記=%d(確認%d) ／ docx置換=%dファイル・%d件 ／ キット差込=%d ／ ルビ付与=%d"
                       % (doc_key, os.path.basename(out_folder),
                          n_x_done, n_x_check, n_files, n_repl, n_kit, n_ruby))

    # 緑ターゲットだが未転記だったセルをレポート出力（要確認）
    if green_report:
        rep = os.path.join(dir_output, "_緑なのに未転記セル一覧.txt")
        try:
            with io.open(rep, "w", encoding="utf-8") as g:
                g.write("緑（転記ターゲット）だが今回転記されなかったセル一覧\n")
                g.write("※ 転記漏れの可能性。不要な緑なら無視可。出力側では黒に戻しています。\n\n")
                for doc_key2, sh, coord, ov in green_report:
                    g.write("[%s] %s!%s  現値=%r\n" % (doc_key2, sh, coord, ov))
            run_log.append("[緑チェック] 未転記の緑セル %d件 → %s" % (len(green_report), os.path.basename(rep)))
        except Exception as e:
            run_log.append("[緑チェック] レポート出力失敗: %r" % e)
    else:
        run_log.append("[緑チェック] 未転記の緑セルなし")
    return out_folders, log_rows


if __name__ == "__main__":
    import datetime
    from transcribe import Hearing, find_hearing, DIR_TPL, DIR_OUTPUT
    hp = find_hearing()
    hearing = Hearing(hp, "ヒアリングシート（PRP）")
    log = []
    outs, _ = run_docs(hearing, hp, DIR_TPL, DIR_OUTPUT, log, datetime.datetime.now())
    print("=== フォルダ型転記 完了 ===")
    print("入力 :", hp)
    for o in outs:
        print("出力 :", o)
    for line in log:
        print("  -", line)
